from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Add image to docx document
def add_image(document, image_path=""):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(5.0))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add title to docx document
def add_title(document, title=""):
    blog_title = document.add_heading(title, 0)
    font = blog_title.style.element.rPr.rFonts
    font.set(qn("w:asciiTheme"), "Arial")
    blog_title.style.font.size = Pt(20)
    blog_title.style.font.color.rgb = RGBColor(0, 0, 0)


# Add subheading to docx document
def add_subheading(document, subheading=""):
    blog_subheading = document.add_heading( subheading, level=1)
    font = blog_subheading.style.element.rPr.rFonts
    font.set(qn("w:asciiTheme"), "Arial")    
    blog_subheading.style.font.color.rgb = RGBColor(0, 0, 0)
    

# Add paragraph to docx document
def add_paragraph(document, paragraph=""):
    paragraph = document.add_paragraph(paragraph)
    paragraph.style.font.name = "Arial"
    paragraph.style.font.size = Pt(12)
    paragraph.style.font.color.rgb = RGBColor(0, 0, 0)